#https://newspaper.readthedocs.io/en/latest/
#API DOCUMENTATIONS FOR EASIER ACCESS OF ARTICLE INFORMATION
#Card cutter - Python Project
#Python 3.4.5 - ATOM Editor
#Ayush Tripathi
import os
from pathlib import Path
import requests

#Import Sc
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
#Import Modules
#Python DOCX Modules
import docx
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

#Import GUI Input Modules
from tkinter import *
from tkinter import simpledialog

#Extra modules
import time
import string
import docx

#Sample Test Wesbites
#https://week.com/2021/03/01/debate-over-criminal-justice-reform-shifts-from-springfield-to-mclean-county/

#Initialize variables
PATH = "C:\Program Files (x86)\chromedriver.exe"

#Take GUI User Input
root = Tk(className='Card Cutter Version 2.1')
root.withdraw()
search_url = simpledialog.askstring(title="Card Cutter Version 2.0", prompt="What article would you like to cut? ")

#Create web driver
options = webdriver.ChromeOptions()
options.add_argument('headless')
options.add_argument('window-size=1920x1080')
options.add_argument("disable-gpu")

driver = webdriver.Chrome(PATH, options=options)
driver.get("https://outline.com")

class Card:
    def __init__(self, url):
        global driver
        self.url = url
        self.author = None
        self.date = None
        self.title = None
        self. content = None
        self.publisher = None

        #Send article into search and hit enter
        search = driver.find_element_by_id("source")
        search.send_keys(self.url)
        search.send_keys(Keys.RETURN)


    def find_author(self):
        global driver
        try:
            author = WebDriverWait(driver, 10).until(
            EC.presence_of_all_elements_located((By.CLASS_NAME, "author")))
            author_raw = author[0].text
            self.author = author_raw.split()[1]
        except:
            author_raw = "NA"
            self.author = "[No Author Detected]"
        return self.author, author_raw


    def find_date(self):
        global driver
        try:
            date = WebDriverWait(driver, 10).until(
            EC.presence_of_all_elements_located((By.CLASS_NAME, "date")))
            date_raw = date[0].text
            self.date = date_raw.split()[2][-2] + date_raw.split()[2][-1]
        except:
            date_raw = "NA"
            self.date = "[No Date Detected]"
        return self.date, date_raw


    def find_title(self):
        global driver
        try:
            title = WebDriverWait(driver, 10).until(
            EC.presence_of_all_elements_located((By.XPATH, 'html/body/outline-app/outline-article/div[2]/div/div[2]/h1')))
            self.title = title[0].text
        except:
            self.title = "Change Title"
        return self.title


    def find_content(self):
        global driver
        try:
            content = WebDriverWait(driver, 10).until(
            EC.presence_of_all_elements_located((By.XPATH, '/html/body/outline-app/outline-article/div[2]/div/raw')))
            self.content = content[0].text
        except:
            self.content = "Evidence Not Found"
        return self.content

#Find article information
c1 = Card(search_url)
c1author, author_raw = c1.find_author()
c1date, date_raw = c1.find_date()
c1title, c1content_raw = c1.find_title(), c1.find_content()
c1content = ''
for i in c1content_raw:
    if i == '\n':
        c1content += ' '
    else:
        c1content += i

#Create site and exit driver
cite = '(' + author_raw + ' ' + search_url + ' ' + date_raw + ')'
driver.quit()

#Python.docx module to convert article information into word document
#Multiple step - local save

try:
    #Create card
    card = Document()

    #Create tagline with correct formatting using title
    tagline = card.add_paragraph()
    tagline.add_run(c1title[0].upper() + c1title[1::].lower()).bold=True
    tagline.style = card.styles.add_style('TAGLINE_STYLE', WD_STYLE_TYPE.PARAGRAPH)
    tagline_font = tagline.style.font
    tagline_font.name = 'Calibri'
    tagline_font.size = Pt(13)

    tagline_format = tagline.paragraph_format
    tagline_format.space_after = Pt(0.5)


    #Create author and cite using local variables
    author = card.add_paragraph()
    author.add_run(c1author[0].upper() + c1author[1::].lower() + ' ' + c1date).bold=True
    author.add_run(' ' + cite).bold=False
    author.style = card.styles.add_style('CITE_STYLE', WD_STYLE_TYPE.PARAGRAPH)
    author_font  = author.style.font
    author_font.name = 'Calibri'
    author_font.size = Pt(11.5)


    #Paste in article content into word document
    content = card.add_paragraph()
    content.add_run(c1content).bold=False
    content.style = card.styles.add_style(('CONTENT_STYLE'), WD_STYLE_TYPE.PARAGRAPH)
    content_style = content.style.font
    content_style.name = 'Calibri'
    content_style.size = Pt(11.75)

    #Save card
    card.save('C:/Users/Desktop/cards/card.docx')

except:
    print("An error has occured. Please restart the program, check to see if the base url is correct, file path is correct, and program has no syntax issues.")
